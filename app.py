import os
import zipfile
import subprocess
import shutil
from io import BytesIO
import pandas as pd
import streamlit as st
from docx import Document

# 1. Configura o visual da página do site
st.set_page_config(page_title="Gerador de Certificados NR", page_icon="🎓", layout="centered")

# Dicionário global para tradução dos meses em português
MESES_PT = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

# Cria as caixas de memória do Streamlit se elas não existirem
if "conectado" not in st.session_state:
    st.session_state["conectado"] = False

# ==========================================================
# SISTEMA DE LOGIN
# ==========================================================
def tela_login():
    st.subheader("🔒 Acesso Restrito")
    usuario = st.text_input("Utilizador", key="user_input")
    senha = st.text_input("Senha", type="password", key="pass_input")
    
    if st.button("Entrar", type="primary"):
        try:
            usuario_correto = st.secrets["credenciais"]["usuario"]
            senha_correta = st.secrets["credenciais"]["senha"]
            
            if usuario == usuario_correto and senha == senha_correta:
                st.session_state["conectado"] = True
                st.success("Acesso concedido! A carregar...")
                st.rerun() 
            else:
                st.error("Utilizador ou senha incorretos.")
        except KeyError:
            st.error("Erro técnico: As credenciais ainda não foram configuradas nos 'Secrets' do Streamlit Cloud.")

# ==========================================================
# DICIONÁRIO DE MODELOS
# ==========================================================
MODELOS = {
    "NR-01 Integração": "modelo_certificadoNR01.docx",
    "NR-06 EPI": "modelo_certificadoNR06.docx",
    "NR-10 Segurança em Eletricidade": "modelo_certificadoNR10.docx",
    "NR-11 Transporte e Movimentação": "modelo_certificadoNR11.docx",
    "NR-12 Máquinas e Equipamentos": "modelo_certificadoNR12.docx",
    "NR-12 Operador de Motosserra": "modelo_certificadoNR12MOTOSSERRA.docx",
    "NR-18 Construção Civil": "modelo_certificadoNR18.docx",
    "NR-23 Combate a Incêndio": "modelo_certificadoNR23.docx",
    "NR-31.7 Segurança na Aplicação de Agrotóxicos": "modelo_certificadoNR31.7.docx",
    "NR-32 Serviços de Saúde": "modelo_certificadoNR32.docx",
    "NR-34 Trabalho a Quente": "modelo_certificadoNR34.docx",
    "NR-35 Trabalho em Altura": "modelo_certificadoNR35.docx"
}

# Lógica avançada para juntar pedaços quebrados de tags pelo Word e manter a formatação
def processar_substituicao(doc, todas_tags):
    def substituir_no_elemento(elemento):
        for paragrafo in elemento.paragraphs:
            for tag, valor in todas_tags.items():
                if tag in paragrafo.text:
                    for i in range(len(paragrafo.runs)):
                        for j in range(i + 1, len(paragrafo.runs) + 1):
                            texto_combinado = "".join([r.text for r in paragrafo.runs[i:j]])
                            if tag in texto_combinado:
                                paragrafo.runs[i].text = texto_combinado.replace(tag, valor)
                                for r in paragrafo.runs[i+1:j]:
                                    r.text = ""
                                break

    substituir_no_elemento(doc)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_no_elemento(celula)

# SE NÃO ESTIVER CONECTADO: Mostra apenas o Login E a Planilha Modelo embaixo
if not st.session_state["conectado"]:
    tela_login()
    st.write("---")
    
    caminho_planilha_modelo = os.path.join("static", "modelo_dados.xlsx")
    if os.path.exists(caminho_planilha_modelo):
        with open(caminho_planilha_modelo, "rb") as f:
            st.download_button(
                label="ℹ️ Descarregar Planilha Modelo de Inserção",
                data=f,
                file_name="modelo_dados.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    st.stop()

# ==========================================================
# PAINEL PRINCIPAL (SÓ APARECE APÓS LOGIN)
# ==========================================================
st.title("🎓 Gerador Web de Certificados NR (PDF)")

if st.sidebar.button("Sair / Logout"):
    st.session_state["conectado"] = False
    st.rerun()

st.write("---")

nr_escolhida = st.selectbox("1. Selecione o Treinamento (NR):", ["Clique para escolher..."] + list(MODELOS.keys()))
arquivo_excel = st.file_uploader("2. Envie a planilha de dados preenchida (.xlsx):", type=["xlsx", "xls"])

if nr_escolhida != "Clique para escolher..." and arquivo_excel is not None:
    if st.button("Processar e Gerar Certificados em PDF", type="primary"):
        try:
            df = pd.read_excel(arquivo_excel)
            
            # Limpa espaços em branco e deixa o cabeçalho em minúsculas para evitar erros de digitação
            df.columns = df.columns.str.strip().str.lower()
            
            caminho_modelo = os.path.join('modelos_docx', MODELOS[nr_escolhida])
            
            if not os.path.exists(caminho_modelo):
                st.error(f"Erro: O modelo '{MODELOS[nr_escolhida]}' não foi encontrado na pasta 'modelos_docx'.")
            else:
                memoria_zip = BytesIO()
                barra_progresso = st.progress(0)
                total_linhas = len(df)
                
                pasta_temp = "temp_certificados"
                os.makedirs(pasta_temp, exist_ok=True)
                
                with zipfile.ZipFile(memoria_zip, 'a', zipfile.ZIP_DEFLATED) as zip_file:
                    for idx, lambda_linha in df.iterrows():
                        doc = Document(caminho_modelo)
                        
                        # CAÇA DA COLUNA DE DATA: Procura qualquer coluna que tenha a palavra "data" no nome
                        coluna_data_real = None
                        for col in df.columns:
                            if "data" in str(col).lower():
                                coluna_data_real = col
                                break
                        
                        data_formatada = ""
                        # Se encontrou alguma coluna de data válida, processa o valor dela
                        if coluna_data_real and pd.notna(lambda_linha.get(coluna_data_real)) and str(lambda_linha[coluna_data_real]).strip() != "":
                            val_data = lambda_linha[coluna_data_real]
                            try:
                                # 1. Se o Pandas leu como um objeto de Data correto do Excel
                                dt = pd.to_datetime(val_data)
                                dia = dt.day
                                ano = dt.year
                                mes_nome = MESES_PT.get(dt.month, "")
                                data_formatada = f"{dia} de {mes_nome} de {ano}"
                            except Exception:
                                try:
                                    # 2. Se o Pandas leu como Texto no formato YYYY-MM-DD
                                    texto_data = str(val_data).split()[0]
                                    partes = texto_data.split('-')
                                    if len(partes) == 3:
                                        ano = int(partes[0])
                                        mes = int(partes[1])
                                        dia = int(partes[2])
                                        mes_nome = MESES_PT.get(mes, "")
                                        data_formatada = f"{dia} de {mes_nome} de {ano}"
                                    else:
                                        # 3. Se o Texto estiver no formato DD/MM/YYYY
                                        partes = texto_data.split('/')
                                        if len(partes) == 3:
                                            dia = int(partes[0])
                                            mes = int(partes[1])
                                            ano = int(partes[2])
                                            mes_nome = MESES_PT.get(mes, "")
                                            data_formatada = f"{dia} de {mes_nome} de {ano}"
                                except Exception:
                                    data_formatada = str(val_data).strip()
                        
                        # Se mesmo assim não encontrou nenhuma coluna com "data", deixa um aviso claro
                        if data_formatada == "":
                            data_formatada = "Coluna de Data não encontrada"
                        
                        # Mapeamento das tags (puxando as colunas em minúsculas limpas pelo script)
                        dados_com_negrito = {
                            "[NOME]": str(lambda_linha.get("nome", "")).strip() if pd.notna(lambda_linha.get("nome")) else "",
                            "[CPF]": str(lambda_linha.get("cpf", "")).strip() if pd.notna(lambda_linha.get("cpf")) else "",
                            "[EMPRESA]": str(lambda_linha.get("empresa", "")).strip() if pd.notna(lambda_linha.get("empresa")) else "",
                            "[CNPJ]": str(lambda_linha.get("cnpj", "")).strip() if pd.notna(lambda_linha.get("cnpj")) else "",
                            "[PERIODO]": str(lambda_linha.get("periodo", "")).strip() if pd.notna(lambda_linha.get("periodo")) else ""
                        }
                        dados_sem_negrito = {"[DATA_FINAL]": data_formatada}
                        todas_tags = {**dados_com_negrito, **dados_sem_negrito}
                        
                        # Executa a substituição avançada mantendo o formato original
                        processar_substituicao(doc, todas_tags)
                                        
                        nome_limpo = str(lambda_linha.get("nome", "aluno")).strip().replace(" ", "_")
                        nr_nome = nr_escolhida.replace(' ', '_')
                        
                        nome_docx = os.path.join(pasta_temp, f"{nr_nome}_{nome_limpo}.docx")
                        nome_pdf = os.path.join(pasta_temp, f"{nr_nome}_{nome_limpo}.pdf")
                        
                        doc.save(nome_docx)
                        
                        # Comando de conversão Pixel Perfect pelo LibreOffice do servidor Linux
                        subprocess.run([
                            'soffice', 
                            '--headless', 
                            '--invisible',
                            '--nodefault',
                            '--nofirststartwizard',
                            '--convert-to', 'pdf:writer_pdf_Export', 
                            '--outdir', pasta_temp, 
                            nome_docx
                        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        
                        if os.path.exists(nome_pdf):
                            zip_file.write(nome_pdf, os.path.basename(nome_pdf))
                        else:
                            zip_file.write(nome_docx, os.path.basename(nome_docx))
                        
                        barra_progresso.progress((idx + 1) / total_linhas)
                
                shutil.rmtree(pasta_temp, ignore_errors=True)
                memoria_zip.seek(0)
                st.success("✨ Todos os certificados foram convertidos e gerados em PDF!")
                
                st.download_button(
                    label="📥 Descarregar Todos os Certificados em PDF (.ZIP)",
                    data=memoria_zip,
                    file_name=f"Certificados_PDF_{nr_escolhida.replace(' ', '_')}.zip",
                    mime="application/zip"
                )
        except Exception as e:
            st.error(f"Ocorreu um erro ao processar a planilha: {e}")

st.write("---")
caminho_planilha_modelo = os.path.join("static", "modelo_dados.xlsx")
if os.path.exists(caminho_planilha_modelo):
    with open(caminho_planilha_modelo, "rb") as f:
        st.download_button(
            label="ℹ️ Descarregar Planilha Modelo de Inserção",
            data=f,
            file_name="modelo_dados.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="modelo_logado"
        )
